"""
objective: this class creates and modifies a word file

"""

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
import os


class new_document:
    def __init__(self, title, file):
        #Create a new document (make self(doc) a document)
        self.titleText = str(title)
        self.fileName = str(file) + ".docx"

        if os.path.exists(self.fileName):
            # open exsisting document
            self.doc = Document(self.fileName)
        else:
            # create new document
            self.doc = Document()
            self.doc.add_heading(self.titleText, 0)

        self.doc.save(self.fileName)

    def newParagraph(self, text):
        "create a new paragraph and add text of a specific colour"
        self.currentParagraph = self.doc.add_paragraph(text)

    def add2Paragraph(self, text):
        "adds text to current paragraph"
        self.currentParagraph.add_run(text)
        
    def addGreenText(self, text):
        "adds green text to current paragraph"
        run = self.currentParagraph.add_run(text)
        run.font.color.rgb = RGBColor(0, 150, 0)

    def addRedText(self, text):
        "adds red text to current paragraph"
        run = self.currentParagraph.add_run(text)
        run.font.color.rgb = RGBColor(150,0,0)

    def addBlueText(self, text):
        "adds blue text to current paragraph"
        run = self.currentParagraph.add_run(text)
        run.font.color.rgb = RGBColor(0,0,150)
    
    def addPageBreak(self):
        "adds a page break"
        self.doc.add_page_break()

    def addHeading(self, heading):
        "adds a heading"
        self.doc.add_heading(str(heading), 1)

    def saveFile(self):
        "save the file"
        self.doc.save(self.fileName)



# print("you have made it to the end of class script")
